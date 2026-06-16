from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "LightGBM" / "User_Manual"
OUT_PATH = OUT_DIR / "LightGBM_V5_User_Manual_TH.docx"


BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(30, 30, 30)
GRAY = RGBColor(90, 90, 90)
LIGHT_BLUE = "EAF2F8"
LIGHT_GREEN = "EAF7EA"
LIGHT_YELLOW = "FFF7D6"
LIGHT_RED = "FDEDEC"
TABLE_HEADER = "D9EAF7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Tahoma"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B8C7D9")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Tahoma"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, RGBColor(48, 84, 150)),
    ]:
        st = styles[name]
        st.font.name = "Tahoma"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        st.paragraph_format.space_after = Pt(5)


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Tahoma"
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.name = "Tahoma"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    else:
        r = p.add_run(text)
        r.font.name = "Tahoma"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Tahoma"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        r.font.size = Pt(10)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        r.font.name = "Tahoma"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        r.font.size = Pt(10)


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Tahoma"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.name = "Tahoma"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    r2.font.size = Pt(10)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=BLUE)
        set_cell_shading(hdr[i], TABLE_HEADER)
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("LightGBM V5 User Manual | Return Risk Prediction Project")
        r.font.name = "Tahoma"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("คู่มือการใช้งานโมเดลทำนายความเสี่ยงการคืนสินค้า")
    r.bold = True
    r.font.name = "Tahoma"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    r.font.size = Pt(20)
    r.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("LightGBM V5 | SETC S1 Clean Dataset 5,000 rows | SETD S3 Real Dataset 55,000 rows")
    r.font.name = "Tahoma"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY

    add_callout(
        doc,
        "เอกสารนี้ใช้ทำอะไร",
        "คู่มือนี้เขียนสำหรับคนที่ต้องรับงานต่อหรือเอาโมเดลไปใช้ต่อ โดยอธิบายแบบภาษาคนว่าโมเดลนี้คืออะไร ใช้ไฟล์ไหนบ้าง ใช้งานอย่างไร ถ้าทำนายพลาดควรตรวจตรงไหน และเมื่อมีข้อมูลจริงของบริษัทเข้ามาควรปรับต่ออย่างไร",
        LIGHT_GREEN,
    )

    doc.add_heading("1. สรุปแบบสั้นที่สุด", level=1)
    add_para(
        doc,
        "โมเดลที่เลือกใช้คือ LightGBM Version 5 หรือ V5 ซึ่งเป็นโมเดลสำหรับทำนายว่า order ใหม่มีโอกาสคืนสินค้าสูงหรือต่ำ โมเดลนี้ไม่ได้ตัดสินแทนคนทั้งหมด แต่ใช้เป็นเครื่องมือช่วยเตือนให้ทีมงานเห็น order ที่ควรตรวจละเอียดขึ้น เช่น โทรยืนยันก่อนส่ง หรือเช็กข้อมูลสินค้าและลูกค้าเพิ่มเติม",
    )
    add_bullets(
        doc,
        [
            "ชุดข้อมูลที่ใช้ train หลัก: SETC/S1 clean_dataset_s1.csv จำนวน 5,000 rows",
            "ชุดข้อมูลที่ใช้ test ภายนอก: SETD/S3 real_dataset_s1.csv จำนวน 55,000 rows",
            "โมเดลที่เลือก: LightGBM V5 เพราะผลทำนายค่อนข้างนิ่ง ใช้ feature ไม่เยอะเกินไป และเหมาะเป็น baseline สำหรับทำระบบจริงต่อ",
            "สถานะโปรเจ็กต์ตอนนี้: ส่วนโมเดลทดลองและเอกสารประกอบพร้อมใช้ต่อแล้ว แต่ถ้าจะใช้กับข้อมูลบริษัทจริง ต้อง map feature และ retrain อีกครั้ง",
        ]
    )

    doc.add_heading("2. โมเดลนี้ช่วยงานอะไร", level=1)
    add_para(
        doc,
        "เมื่อมี order ใหม่เข้ามา ระบบจะดูข้อมูลลูกค้า สินค้า หมวดสินค้า ช่องทางชำระเงิน ช่องทางขาย ขนส่ง และประวัติการซื้อ/คืนย้อนหลัง แล้วให้คะแนนความเสี่ยงว่า order นี้มีโอกาสคืนสินค้ามากน้อยแค่ไหน",
    )
    add_table(
        doc,
        ["ผลจากโมเดล", "ความหมาย", "ตัวอย่างการนำไปใช้"],
        [
            ["Low Risk", "โอกาสคืนสินค้าต่ำ", "ดำเนินการตามขั้นตอนปกติ"],
            ["Medium Risk", "มีความเสี่ยงระดับกลาง", "เช็กข้อมูลสินค้า/ที่อยู่/การชำระเงินเพิ่มเติม"],
            ["High Risk", "มีโอกาสคืนสูงกว่าปกติ", "โทรยืนยันก่อนจัดส่ง หรือให้ทีมงานตรวจ order ก่อน"],
        ],
        [1.4, 2.2, 3.0],
    )

    doc.add_heading("3. ทำไมถึงเลือก LightGBM V5", level=1)
    add_para(
        doc,
        "V5 เป็นเวอร์ชันที่เลือก feature แบบพอดี ไม่เบาเกินไปแบบ V1 และไม่หนักเกินไปแบบ V4 จุดเด่นคือยังเก็บ feature สำคัญ เช่น ประวัติลูกค้า ความเสี่ยงสินค้า ความเสี่ยงหมวดสินค้า วิธีชำระเงิน และ interaction บางตัวไว้ แต่ตัด feature ที่ทำให้ระบบหนักหรือซ้ำซ้อนออก",
    )
    add_table(
        doc,
        ["ชุดทดสอบ", "จำนวนข้อมูล", "Accuracy", "Recall", "F1", "AUC", "หมายเหตุ"],
        [
            ["SETC S1 Holdout", "1,000 rows จาก clean 5,000", "82.00%", "59.56%", "67.86%", "83.27%", "ใช้วัดหลัง train โดยแยก holdout"],
            ["SETD S3 External", "55,000 rows", "81.91%", "61.24%", "68.37%", "82.59%", "ใช้ดูว่าโมเดลเจอข้อมูลใหม่แล้วยังนิ่งไหม"],
        ],
        [1.35, 1.25, 0.85, 0.85, 0.85, 0.85, 1.6],
    )
    add_callout(
        doc,
        "ข้อควรจำ",
        "Accuracy 82% ไม่ได้แปลว่าใช้จริงแล้วจะได้ 82% เสมอไป เพราะข้อมูลจริงของบริษัทอาจมี pattern ต่างจากข้อมูลที่ใช้ทดลอง เมื่อมีข้อมูลจริงต้องนำมาทดสอบและ retrain อีกครั้ง",
        LIGHT_YELLOW,
    )

    doc.add_heading("4. ไฟล์ที่ต้องใช้ถ้าจะเอาโมเดลไปใช้ต่อ", level=1)
    add_para(doc, "ถ้าคนอื่นจะรับงานต่อ ให้ดูไฟล์ชุดนี้ก่อน เพราะเป็นไฟล์หลักของ V5 ที่เลือกไว้")
    add_table(
        doc,
        ["ไฟล์", "ใช้ทำอะไร"],
        [
            ["docs/LightGBM/SETC/clean_dataset/S1/V5/models/model_lgbm_s1_v5_lightgbm.pkl", "ไฟล์โมเดลจริง ใช้สำหรับทำนายผล"],
            ["docs/LightGBM/SETC/clean_dataset/S1/V5/models/model_lgbm_s1_v5_metadata.json", "บอกค่าที่ใช้กับโมเดล เช่น threshold, parameter, accuracy"],
            ["docs/LightGBM/SETC/clean_dataset/S1/V5/features/used_features_lgbm_s1_v5.csv", "รายชื่อ feature 64 ตัวที่ต้องป้อนเข้าโมเดล"],
            ["docs/LightGBM/SETC/clean_dataset/S1/V5/features/train_validation_holdout_sets_lgbm_s1_v5.pkl", "ข้อมูล split และ schema ที่ใช้ตอน train/test"],
            ["docs/LightGBM/SETC/clean_dataset/clean_dataset_s1.csv", "clean dataset 5,000 rows ที่ใช้เป็นฐาน train"],
            ["docs/LightGBM/SETD/real_dataset/S3/real_dataset_s1.csv", "real-like dataset 55,000 rows ที่ใช้ test ทั้งก้อน"],
            ["docs/LightGBM/SETD/real_dataset/S3/V5/reports/external_metrics_lgbm_s1_v5.csv", "ผล test ของ V5 บน real dataset S3"],
        ],
        [3.5, 3.0],
    )

    doc.add_heading("5. วิธีเอาโมเดลไปใช้ต่อแบบภาษาคน", level=1)
    add_para(
        doc,
        "เวลาจะนำโมเดลไปใช้กับระบบจริง ไม่ใช่เอาไฟล์ CSV ทั้งหมดโยนเข้าโมเดลทุกครั้ง แต่จะใช้กับ order ใหม่ทีละรายการ หรือทีละชุด order ที่เข้ามา ระบบหลังบ้านจะเตรียมข้อมูลให้เหมือนตอน train แล้วจึงให้โมเดลทำนาย",
    )
    add_numbered(
        doc,
        [
            "รับ order ใหม่ เช่น customer_id, product_id, ช่องทางขาย, วิธีชำระเงิน, จำนวนสินค้า และราคา",
            "ดึงประวัติของลูกค้าคนนี้เท่านั้น ไม่ต้องดึงข้อมูลลูกค้าทั้งระบบ",
            "ดึงข้อมูลสินค้า หมวดสินค้า แบรนด์ ขนส่ง และค่าสถิติเกี่ยวกับการคืนสินค้า",
            "สร้าง feature 1 row ให้มีหน้าตาเหมือน used_features_lgbm_s1_v5.csv",
            "ส่ง feature 1 row เข้าโมเดล LightGBM V5",
            "โมเดลตอบกลับเป็นความเสี่ยง เช่น 0.72 หรือ 72%",
            "ระบบแปลงคะแนนเป็น Low / Medium / High Risk แล้วแสดงให้ผู้ใช้งาน",
        ]
    )
    add_table(
        doc,
        ["ตัวอย่างสิ่งที่ผู้ใช้งานเห็น", "คำอธิบาย"],
        [
            ["Risk Probability = 72%", "order นี้มีความเสี่ยงคืนสินค้าค่อนข้างสูง"],
            ["Risk Level = High Risk", "ควรให้พนักงานตรวจเพิ่มก่อนส่ง"],
            ["Action = โทรยืนยันก่อนจัดส่ง", "ลดโอกาสส่งผิด/ลูกค้าไม่รับ/คืนสินค้า"],
        ],
        [2.3, 4.2],
    )

    doc.add_heading("6. Feature ที่โมเดล V5 ใช้คืออะไร", level=1)
    add_para(
        doc,
        "Feature คือข้อมูลที่ถูกแปลงให้อยู่ในรูปที่โมเดลเข้าใจได้ V5 ใช้ 64 feature แต่สามารถแบ่งเป็นกลุ่มใหญ่ๆ ได้ดังนี้",
    )
    add_table(
        doc,
        ["กลุ่ม feature", "ตัวอย่าง", "ความหมายแบบง่าย"],
        [
            ["ข้อมูลลูกค้า", "age, membership_tier, province", "ลูกค้าอยู่กลุ่มไหน อายุเท่าไหร่ อยู่จังหวัดใด"],
            ["ข้อมูลสินค้า/order", "category, brand, quantity, total_amount", "ซื้อสินค้าอะไร ราคาเท่าไหร่ จำนวนเท่าไหร่"],
            ["ประวัติลูกค้า", "total_orders_before, total_returns_before, customer_return_ratio", "ลูกค้าคนนี้เคยซื้อและเคยคืนบ่อยแค่ไหน"],
            ["ประวัติย้อนหลังตามช่วงเวลา", "hist_return_rate_7d, 30d, 90d, 365d", "ช่วงหลังๆ ลูกค้าคืนบ่อยขึ้นหรือลดลง"],
            ["ความเสี่ยงสินค้า/หมวด", "product_return_rate_pti, category_return_rate_pti, brand_return_rate_pti", "สินค้าหรือหมวดนี้เคยมีการคืนสูงไหม"],
            ["ความเสี่ยงขนส่ง/ช่องทาง", "courier_return_rate_pti, payment_return_rate_pti, channel_return_rate_pti", "ขนส่ง วิธีจ่ายเงิน หรือช่องทางขายนี้มี pattern คืนสินค้าไหม"],
            ["feature ผสม", "category_payment, high_discount_cod, low_rating_high_discount", "จับคู่ปัจจัย เช่น หมวดสินค้า + COD หรือส่วนลดสูง + rating ต่ำ"],
        ],
        [1.6, 2.2, 2.8],
    )

    doc.add_heading("7. ถ้าโมเดลทำนายพลาด ต้องแก้ตรงไหน", level=1)
    add_para(
        doc,
        "การทำนายพลาดเกิดได้ปกติ เพราะโมเดลเรียนรู้จาก pattern ในอดีต ไม่ได้รู้ความจริงล่วงหน้า วิธีแก้ต้องดูว่าพลาดแบบไหน",
    )
    add_table(
        doc,
        ["ปัญหา", "ความหมาย", "ควรตรวจ/แก้ตรงไหน"],
        [
            ["โมเดลบอกเสี่ยงสูง แต่ลูกค้าไม่คืน", "False Positive", "threshold อาจต่ำเกินไป, feature บางตัวให้ความเสี่ยงสูงเกินจริง, กลุ่มลูกค้า/สินค้าอาจเปลี่ยน pattern"],
            ["โมเดลบอกเสี่ยงต่ำ แต่ลูกค้าคืนจริง", "False Negative", "โมเดลจับเคสเสี่ยงไม่ทัน, อาจต้องเพิ่ม feature เช่น return reason, product defect, call center note หรือปรับ threshold ให้เข้มขึ้น"],
            ["ข้อมูลใหม่มี category/courier/payment ที่ไม่เคยเจอ", "ข้อมูลใหม่ไม่เหมือนตอน train", "ต้อง update mapping, เพิ่มค่าใหม่ใน feature builder และ retrain"],
            ["Accuracy ตกเมื่อใช้ข้อมูลจริง", "ข้อมูลจริงต่างจากข้อมูลทดลอง", "ทำ EDA เทียบ distribution, ตรวจ missing/outlier, retrain ด้วยข้อมูลจริง"],
            ["ระบบทำนายช้า", "feature engineering ใช้เวลานาน", "ใช้ Feature Store/cache เก็บค่าประวัติลูกค้าและ return rate ไว้ล่วงหน้า"],
        ],
        [1.35, 1.65, 3.55],
    )
    add_callout(
        doc,
        "จุดที่ควรแก้ก่อนเมื่อทำนายพลาดบ่อย",
        "เริ่มจากตรวจข้อมูลที่ป้อนเข้าโมเดลก่อนเสมอ เช่น customer_id ถูกไหม category ตรงไหม payment_method ตรงไหม และ feature ประวัติลูกค้าคำนวณจากข้อมูลก่อน order ปัจจุบันจริงหรือไม่ หลังจากนั้นค่อยปรับ threshold หรือ retrain",
        LIGHT_RED,
    )

    doc.add_heading("8. เมื่อมีข้อมูลจริงของบริษัท ต้องทำอะไรบ้าง", level=1)
    add_para(
        doc,
        "เมื่อบริษัทให้ข้อมูลจริงมาแล้ว ไม่ควรเอาโมเดลเดิมไปใช้ทันทีแบบไม่ตรวจ เพราะข้อมูลจริงอาจมีชื่อ column, รูปแบบข้อมูล, สัดส่วน return, หมวดสินค้า หรือพฤติกรรมลูกค้าต่างจาก dataset ทดลอง",
    )
    add_numbered(
        doc,
        [
            "รับข้อมูลจริงจากบริษัท เช่น order, customer, product, return, courier, payment และ promotion",
            "ทำ Data Dictionary ว่าแต่ละ column หมายถึงอะไร",
            "Map column จริงเข้ากับ feature ที่ V5 ต้องใช้ เช่น customer_id, order_date, category, payment_method",
            "ตรวจ missing/null/outlier และ duplicate",
            "สร้าง feature ด้วยสูตรเดียวกับ V5 โดยใช้ข้อมูลก่อน order ปัจจุบันเท่านั้น",
            "ทดสอบ V5 เดิมกับข้อมูลจริงเพื่อดู baseline",
            "Retrain โมเดลด้วยข้อมูลจริง แล้วเทียบกับ V5 เดิม",
            "Tune threshold/parameter ให้เหมาะกับต้นทุนจริงของธุรกิจ",
            "เลือกโมเดลสุดท้ายอีกครั้งก่อนขึ้นระบบจริง",
        ]
    )

    doc.add_heading("9. ต้อง retrain หรือไม่", level=1)
    add_para(
        doc,
        "ต้อง retrain เมื่อมีข้อมูลจริงครับ เพราะตอนนี้โมเดลเรียนรู้จากข้อมูลทดลอง/เสมือนจริงที่ออกแบบให้ใกล้ธุรกิจ O Shopping แต่ยังไม่ใช่ production data จริงทั้งหมด",
    )
    add_table(
        doc,
        ["สถานการณ์", "ต้อง retrain ไหม", "เหตุผล"],
        [
            ["ข้อมูลจริงมี schema เหมือนเดิม และ pattern ใกล้เดิม", "ควร retrain แต่ไม่เร่งด่วน", "ใช้ V5 เดิมเป็น baseline ได้ก่อน"],
            ["return rate จริงต่างจาก dataset ทดลองมาก", "ต้อง retrain", "โมเดลจะ bias ตามสัดส่วนเดิม"],
            ["มีหมวดสินค้า/ขนส่ง/ช่องทางใหม่", "ต้อง retrain หรืออย่างน้อย update mapping", "โมเดลไม่รู้จัก pattern ใหม่"],
            ["มี feature ใหม่ เช่น call center note หรือเหตุผลคืนสินค้า", "ควรสร้าง V6", "ข้อมูลใหม่อาจช่วยเพิ่มความแม่นยำ"],
            ["โมเดลทำนายผิดซ้ำๆ ในบางกลุ่มสินค้า", "ต้อง retrain เฉพาะรอบ", "pattern ของสินค้ากลุ่มนั้นอาจเปลี่ยน"],
        ],
        [2.0, 1.35, 3.15],
    )

    doc.add_heading("10. โปรเจ็กต์นี้สมบูรณ์หรือยัง", level=1)
    add_callout(
        doc,
        "คำตอบตรงๆ",
        "ส่วนโมเดลและการทดลองถือว่าพร้อมเป็น baseline สำหรับส่งต่อและพัฒนาต่อ แต่ยังไม่ใช่ระบบ production ที่สมบูรณ์ 100% เพราะยังต้องต่อ API, Feature Store, ระบบหน้าเว็บ, monitoring และ retraining pipeline เมื่อมีข้อมูลจริง",
        LIGHT_YELLOW,
    )
    add_table(
        doc,
        ["ส่วนงาน", "สถานะ", "สิ่งที่ยังต้องทำต่อ"],
        [
            ["Dataset ทดลอง SETC S1", "พร้อมใช้", "ใช้เป็น baseline เท่านั้น ไม่แทนข้อมูลจริง"],
            ["Model LightGBM V5", "พร้อมส่งต่อ", "ต้อง retrain เมื่อมีข้อมูลจริง"],
            ["Feature Engineering", "พร้อมเป็นแนวทาง", "ต้อง map กับ column จริงของบริษัท"],
            ["External Test SETD S3", "มีผลทดสอบแล้ว", "ควรมี real production holdout แยกอีกชุดเมื่อได้ข้อมูลจริง"],
            ["Web/API", "ยังไม่ใช่ production เต็ม", "ต้องทำ FastAPI/Frontend/Database/Feature Store"],
            ["Monitoring", "ยังต้องเพิ่ม", "ต้องเก็บผลทำนายกับผลคืนจริงเพื่อวัด drift"],
        ],
        [1.65, 1.25, 3.65],
    )

    doc.add_heading("11. คู่มือการใช้งานสำหรับคนที่ไม่ใช่สายคอม", level=1)
    add_para(doc, "ถ้าใช้งานผ่านเว็บในอนาคต ผู้ใช้ทั่วไปไม่จำเป็นต้องเปิดไฟล์โมเดลเอง ให้ทำตามขั้นตอนนี้")
    add_numbered(
        doc,
        [
            "เปิดหน้าเว็บ Return Risk Prediction",
            "ใส่รหัสลูกค้า หรือค้นหาลูกค้าจาก order ที่เข้ามา",
            "ระบบจะแสดงประวัติการซื้อ/คืนย้อนหลังของลูกค้า",
            "ตรวจว่าข้อมูล order ใหม่ถูกต้อง เช่น สินค้า จำนวน ราคา วิธีจ่ายเงิน",
            "กดปุ่ม Predict หรือ ประเมินความเสี่ยง",
            "ดูผลลัพธ์ว่าเป็น Low, Medium หรือ High Risk",
            "ถ้า High Risk ให้ทำ action ตามนโยบาย เช่น โทรยืนยันก่อนจัดส่ง",
            "หลัง order จบจริง ให้บันทึกผลว่า return หรือ not return เพื่อใช้ปรับโมเดลในอนาคต",
        ]
    )

    doc.add_heading("12. Production Flow ที่ควรทำจริง", level=1)
    add_para(
        doc,
        "ระบบจริงควรทำแบบนี้ เพื่อไม่ให้ช้าและไม่ต้องคำนวณทั้ง dataset ทุกครั้ง",
    )
    add_table(
        doc,
        ["ขั้นตอน", "สิ่งที่เกิดขึ้น"],
        [
            ["1. Order ใหม่เข้ามา", "ระบบรับข้อมูล order ใหม่ 1 รายการ"],
            ["2. Query เฉพาะลูกค้าคนนี้", "ดึงประวัติเฉพาะ customer_id ที่เกี่ยวข้อง"],
            ["3. Query snapshot ที่เกี่ยวข้อง", "ดึง product/category/courier/payment/channel summary"],
            ["4. Build feature 1 row", "สร้างข้อมูลให้ตรงกับ used_features_lgbm_s1_v5.csv"],
            ["5. Predict", "ส่งเข้า LightGBM V5 และรับคะแนนความเสี่ยง"],
            ["6. แสดงผล", "แสดง risk level และ action ที่ควรทำ"],
            ["7. เก็บ feedback", "เมื่อรู้ผลคืนจริง ให้เก็บกลับเพื่อ retrain รอบถัดไป"],
        ],
        [1.8, 4.7],
    )

    doc.add_heading("13. Checklist ก่อนส่งต่อให้คนอื่นใช้", level=1)
    add_bullets(
        doc,
        [
            "มีไฟล์โมเดล .pkl อยู่ครบ",
            "มีไฟล์ used_features_lgbm_s1_v5.csv เพื่อรู้ว่าโมเดลต้องใช้ feature อะไร",
            "มี metadata.json เพื่อดู threshold และ parameter",
            "มี script หรือระบบที่สร้าง feature ได้เหมือนตอน train",
            "ข้อมูลใหม่ต้องมี column ที่ map กับ feature ได้",
            "ห้ามใช้ข้อมูลหลังเหตุการณ์ เช่น return_date ของ order ปัจจุบัน, refund_amount หรือ return_reason ตอนทำนาย order ใหม่",
            "ต้องเก็บผล predict และผลจริงไว้ตรวจย้อนหลัง",
            "ถ้า accuracy ลดลง ต้องเช็กข้อมูลก่อน แล้วค่อย retrain/tune",
        ]
    )

    doc.add_heading("14. คำอธิบายคำศัพท์", level=1)
    add_table(
        doc,
        ["คำ", "ความหมายแบบง่าย"],
        [
            ["Model", "ตัวช่วยทำนายจากข้อมูลในอดีต"],
            ["Feature", "ข้อมูลที่แปลงแล้วเพื่อให้โมเดลใช้คิด"],
            ["Train", "ให้โมเดลเรียนรู้จากข้อมูลที่มีคำตอบแล้ว"],
            ["Test", "นำโมเดลไปสอบกับข้อมูลอีกชุดเพื่อดูว่าทายแม่นแค่ไหน"],
            ["Accuracy", "ทายถูกทั้งหมดกี่เปอร์เซ็นต์"],
            ["Recall", "จับเคสที่คืนสินค้าได้มากแค่ไหน"],
            ["Threshold", "เส้นแบ่งว่าจะนับว่าเสี่ยงหรือไม่เสี่ยง"],
            ["Retrain", "ฝึกโมเดลใหม่ด้วยข้อมูลชุดใหม่"],
            ["Feature Store", "ที่เก็บ feature ที่คำนวณไว้ล่วงหน้า เพื่อให้ระบบทำนายเร็วขึ้น"],
        ],
        [1.45, 5.05],
    )

    doc.add_heading("15. สรุปสำหรับส่งต่องาน", level=1)
    add_para(
        doc,
        "ตอนนี้งานที่ส่งต่อได้คือโมเดล LightGBM V5 พร้อมไฟล์ประกอบสำหรับ train/test และผลประเมินบน SETC S1 กับ SETD S3 โมเดลนี้เหมาะใช้เป็น baseline สำหรับระบบทำนายความเสี่ยงการคืนสินค้า แต่เมื่อมีข้อมูลจริงของบริษัทเข้ามา ต้อง map feature, ตรวจคุณภาพข้อมูล, retrain, tune threshold และทดสอบอีกครั้งก่อนใช้จริงใน production",
    )
    add_callout(
        doc,
        "ประโยคสั้นๆ สำหรับอธิบายคนรับงานต่อ",
        "เราเลือก LightGBM V5 เพราะเป็นรุ่นที่สมดุลที่สุดระหว่างความแม่นยำ ความนิ่งเมื่อเจอข้อมูลใหม่ และจำนวน feature ที่ไม่หนักเกินไป ตอนนี้ใช้ต่อได้เป็น baseline แต่ถ้าจะใช้กับข้อมูลบริษัทจริง ต้อง retrain และตรวจ feature mapping ก่อนขึ้นระบบจริง",
        LIGHT_GREEN,
    )

    add_footer(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
